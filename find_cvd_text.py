import docx
import re

def find_cvd_text(filename):
    doc = docx.Document(filename)
    
    keywords = ['cardiovascular', 'heart', 'disease', 'patient', 'clinical', 'medical', 'blood', 'cvd', 'accuracy', 'predict']
    # 'accuracy' and 'predict' are generic but might help find residual ML text that hasn't been adapted.
    # Let's stick to more specific ones first to avoid false positives.
    specific_keywords = ['cardiovascular', 'heart ', 'disease', 'patient', 'clinical', 'medical', 'blood ', 'cvd', 'vessel', 'cholesterol', 'sugar', 'electrocardiographic', 'angina']
    
    print(f"Scanning {filename} for sample template remnants...")
    
    found_any = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        # check if any specific keyword is in text
        matches = [kw for kw in specific_keywords if kw in text]
        if matches:
            print(f"\n--- Paragraph {i} ---")
            print(f"Contains: {', '.join(matches)}")
            print(f"Text:\n{para.text}")
            found_any = True
            
    # Also check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.lower()
                    matches = [kw for kw in specific_keywords if kw in text]
                    if matches:
                        print(f"\n--- Table {t_idx}, Row {r_idx}, Col {c_idx}, Para {p_idx} ---")
                        print(f"Contains: {', '.join(matches)}")
                        print(f"Text:\n{para.text}")
                        found_any = True

    if not found_any:
        print("No matches found for CVD keywords.")

if __name__ == '__main__':
    find_cvd_text('AEGIS_EARTH_DOCS_v2.docx')
